"""
降 AIGC 率工作流
流程：检测 → 识别高风险段落 → 分类 → 逐段润色 → 复测验证

支持三种润色策略：
1. 表达润色（中文论文）— 修复语病、去除口语，克制修改
2. 逻辑检查 — 红线审查，仅报致命错误
3. 去 AI 味 — 将机械化文本重写为自然学术表达
"""
import os
import re
import json
import logging
import time
import asyncio
from datetime import datetime
from typing import Callable, Dict, List, Optional, Tuple

from .progress import progress_bar

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════
# Prompt 模板
# ═══════════════════════════════════════════════════════════

PROMPT_REFINE_CN = """你是一位专注于计算机科学领域的资深中文学术编辑，深谙《计算机学报》、《软件学报》等核心期刊的审稿标准。你秉持尊重原著、克制修改的原则，具备敏锐的鉴赏力，只在确有必要时才进行干预。

请对以下中文论文段落进行专业审视与润色。核心任务：修复明显的语病与逻辑漏洞。如果原文表达已经清晰、准确且符合学术规范，请务必保留原样。

约束：
1. 修正阈值：仅在检测到口语化表达（如"我们觉得"）、语法错误、逻辑断层或严重欧化长句时才修正。禁止为追求形式变化而强行替换同义词或重组句式。
2. 语体规范：坚持当代学术书面语，平实、流畅、准确。禁止将"旨在"改为"拟"，将"是"改为"系"。彻底去除口语。
3. 逻辑与连贯性：仅在逻辑断裂时显化连接词，否则优先依赖语序自然衔接。
4. 格式：输出纯文本，禁止 Markdown。严格使用中文全角标点。

输出要求：
- 第一行输出润色后的文本（如无需修改则原样输出）
- 第二行起简要说明修改点（如无需修改则输出"原文表达规范，未做修改。"）

原文：
{paragraph}"""


PROMPT_DEAI_CN = """你是一位计算机科学领域的资深学术编辑，专注于提升中文论文的自然度与可读性。你的任务是将大模型生成的机械化中文文本重写为自然的学术表达，降低 AIGC 检测率。

重写原则：
1. 词汇去模板化：消除"综上所述"、"值得注意的是"、"从本质上讲"等 AI 高频套话，用更自然的学术表达替代。
2. 句式多样化：避免连续使用相同句式结构（如多个"首先...其次...最后..."），混合使用长短句。
3. 增加个人化痕迹：适当加入"本文认为"、"实验发现"、"在调试过程中"等体现作者主体性的表达。
4. 细节具体化：将泛泛而谈的概述替换为具体的技术细节、参数、代码片段等。
5. 保留技术准确性：不得改变原文的技术含义和数据。
6. 格式：输出纯文本，禁止 Markdown。严格使用中文全角标点。

直接输出重写后的文本，不要输出任何解释。

原文：
{paragraph}"""


PROMPT_DEAI_EN = """You are a senior academic editor in computer science. Rewrite the following English academic text to sound more natural and human-written, reducing AI-generated patterns.

Rules:
1. Prefer plain, precise academic vocabulary. Avoid overused words like "leverage", "delve into", "tapestry". Use "use", "investigate", "context" instead.
2. Do NOT use list format. Convert all bullet points into flowing prose paragraphs.
3. Remove mechanical transition phrases ("First and foremost", "It is worth noting that"). Connect ideas through logical progression.
4. Vary sentence length and structure. Avoid repetitive patterns.
5. Preserve all technical accuracy and data.
6. Output plain text only, no Markdown formatting.

Rewritten text:

{paragraph}"""


PROMPT_LOGIC_CHECK_EN = """You are a proofreader for final drafts. Perform a "red-line review" on the following English LaTeX code snippet. Only flag fatal errors.

Default assumption: the draft has been through multiple rounds of revision and is high quality. Only report issues that obstruct reader understanding.

Review dimensions:
- Fatal logic: completely contradictory statements?
- Term consistency: has a core concept changed names without explanation?
- Severe grammar errors: Chinglish or structural errors that obscure meaning.

If there are no "must-fix" errors, output exactly: [检测通过，无实质性问题]
If there are issues, list them briefly in Chinese.

Input:
{paragraph}"""


PROMPTS = {
    'refine_cn': PROMPT_REFINE_CN,
    'deai_cn': PROMPT_DEAI_CN,
    'deai_en': PROMPT_DEAI_EN,
    'logic_check_en': PROMPT_LOGIC_CHECK_EN,
}


PROMPT_ITERATIVE_HINT_CN = """上一轮润色后的 AIGC 分数仍然偏高（{score:.1f}%），请在不改变技术含义、数据、术语、公式和引用的前提下，进一步降低模板化表达。

要求：
1. 避免机械连接词和套话。
2. 保留原有信息密度，不要扩写空话。
3. 句式可以更自然，但不得引入新事实。
4. 直接输出改写后的纯文本，不要解释。

待调整文本：
{paragraph}"""


# ═══════════════════════════════════════════════════════════
# 段落分类器
# ═══════════════════════════════════════════════════════════

def classify_paragraph(text: str) -> str:
    """
    分类段落类型，选择最合适的润色策略
    Returns: 'refine_cn' / 'deai_cn' / 'deai_en' / 'skip'
    """
    text = text.strip()

    # 跳过过短的段落
    if len(text) < 30:
        return 'skip'

    # 跳过代码块、公式
    if text.startswith('```') or text.startswith('$') or re.match(r'^[A-Z_]+\s*[=]', text):
        return 'skip'

    # 跳过参考文献条目
    if re.match(r'^\[\d+\]', text):
        return 'skip'

    # 跳过图表标题
    if re.match(r'^[图表]\d', text):
        return 'skip'

    # 检测是否为英文
    ascii_chars = sum(1 for c in text if ord(c) < 128 and c.isalpha())
    total_alpha = sum(1 for c in text if c.isalpha())
    if total_alpha > 0 and ascii_chars / total_alpha > 0.5:
        return 'deai_en'

    # 中文段落：根据 AI 特征程度选择策略
    # 如果含有高密度 AI 套话，用 deai_cn
    ai_markers = [
        '综上所述', '值得注意的是', '从本质上讲', '换句话说',
        '具体来说', '需要指出的是', '不难发现', '毋庸置疑',
        '首先，', '其次，', '最后，', '一方面', '另一方面',
    ]
    marker_count = sum(1 for m in ai_markers if m in text)
    if marker_count >= 2:
        return 'deai_cn'

    return 'refine_cn'


# ═══════════════════════════════════════════════════════════
# 润色引擎
# ═══════════════════════════════════════════════════════════

class RefinementEngine:
    """
    降 AIGC 率润色引擎

    使用 LLM API 对高风险段落进行智能润色
    """

    def __init__(
        self,
        api_base: str = "https://api.openai.com/v1",
        api_key: str = "",
        model: str = "gpt-4o-mini",
        temperature: float = 0.3,
        max_retries: int = 3,
        concurrency: int = 10,
    ):
        self.api_base = api_base.rstrip('/')
        self.api_key = api_key or os.environ.get("OPENAI_API_KEY", "")
        self.model = model
        self.temperature = temperature
        self.max_retries = max_retries
        self.concurrency = concurrency

    def _get_client(self):
        from openai import OpenAI
        return OpenAI(api_key=self.api_key, base_url=self.api_base, timeout=120, max_retries=self.max_retries)

    def _get_async_client(self):
        from openai import AsyncOpenAI
        return AsyncOpenAI(api_key=self.api_key, base_url=self.api_base, timeout=120, max_retries=self.max_retries)

    @staticmethod
    def _parse_refinement_result(text: str, strategy: str) -> Tuple[str, str]:
        if strategy == 'logic_check_en':
            return text, "逻辑检查不修改原文"
        if strategy == 'refine_cn':
            lines = text.split('\n', 1)
            refined = lines[0].strip()
            note = lines[1].strip() if len(lines) > 1 else "已润色"
            if refined.startswith('"') and refined.endswith('"'):
                refined = refined[1:-1]
            if refined.startswith('"') and refined.endswith('"'):
                refined = refined[1:-1]
            return refined, note
        return text.strip(), "已重写（去 AI 化）"

    def _call_llm(self, prompt: str) -> Optional[str]:
        """调用 LLM API"""
        try:
            client = self._get_client()
            resp = client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=self.temperature,
                max_tokens=2000,
            )
            return resp.choices[0].message.content.strip()
        except Exception as e:
            logger.warning("API 调用失败: %s", e)
            return None

    def refine_paragraph(self, text: str, strategy: str) -> Tuple[str, str]:
        """
        润色单个段落

        Args:
            text: 原文段落
            strategy: 润色策略 ('refine_cn' / 'deai_cn' / 'deai_en' / 'logic_check_en')

        Returns:
            (refined_text, change_note)
        """
        if strategy == 'skip':
            return text, "跳过（过短/代码/公式）"

        prompts = {
            'refine_cn': PROMPT_REFINE_CN,
            'deai_cn': PROMPT_DEAI_CN,
            'deai_en': PROMPT_DEAI_EN,
            'logic_check_en': PROMPT_LOGIC_CHECK_EN,
        }

        prompt_template = prompts.get(strategy, PROMPT_DEAI_CN)
        prompt = prompt_template.format(paragraph=text)

        result = self._call_llm(prompt)
        if not result:
            return text, "API 调用失败，保留原文"

        if strategy == 'logic_check_en':
            return text, result  # 逻辑检查不改原文，只输出检查结果

        return self._parse_refinement_result(result, strategy)

    async def _call_llm_async(self, prompt: str, client) -> Optional[str]:
        """异步调用 LLM API"""
        resp = await client.chat.completions.create(
            model=self.model,
            messages=[{"role": "user", "content": prompt}],
            temperature=self.temperature,
            max_tokens=2000,
        )
        return resp.choices[0].message.content.strip()

    async def _refine_one_async(self, para: Dict, score_info: Dict,
                                score_threshold: float, semaphore: asyncio.Semaphore,
                                client) -> Dict:
        """异步润色单个段落（带并发控制）"""
        text = para['text']
        aigc_score = score_info.get('aigc_score')

        # 低于阈值跳过
        if aigc_score is not None and aigc_score <= score_threshold:
            return {
                'index': para['index'],
                'chapter': para.get('chapter', 0),
                'original': text,
                'refined': text,
                'strategy': 'skip',
                'changed': False,
                'note': f'AIGC={aigc_score*100:.1f}%，低于阈值，跳过',
                'aigc_before': aigc_score,
                'aigc_after': None,
            }

        strategy = classify_paragraph(text)

        if strategy == 'skip':
            return {
                'index': para['index'],
                'chapter': para.get('chapter', 0),
                'original': text,
                'refined': text,
                'strategy': 'skip',
                'changed': False,
                'note': '跳过（代码/公式/过短）',
                'aigc_before': aigc_score,
                'aigc_after': None,
            }

        try:
            async with semaphore:
                result = await self._call_llm_async(
                    PROMPTS[strategy].format(paragraph=text), client
                )
        except Exception as e:
            logger.warning("段落 %d API 调用失败: %s", para['index'], e)
            result = None

        if not result:
            return {
                'index': para['index'],
                'chapter': para.get('chapter', 0),
                'original': text,
                'refined': text,
                'strategy': strategy,
                'changed': False,
                'note': 'API 调用失败，保留原文',
                'aigc_before': aigc_score,
                'aigc_after': None,
            }

        if strategy == 'logic_check_en':
            refined, note = text, result
        else:
            refined, note = self._parse_refinement_result(result, strategy)

        changed = (refined.strip() != text.strip())
        return {
            'index': para['index'],
            'chapter': para.get('chapter', 0),
            'original': text,
            'refined': refined,
            'strategy': strategy,
            'changed': changed,
            'note': note,
            'aigc_before': aigc_score,
            'aigc_after': None,
        }

    async def _detect_async(
        self,
        text: str,
        detect_fn: Callable[[str], Dict],
        semaphore: asyncio.Semaphore,
    ) -> Dict:
        async with semaphore:
            return await asyncio.to_thread(detect_fn, text)

    @staticmethod
    async def _with_index(index: int, coro):
        try:
            return index, await coro
        except Exception as e:
            return index, e

    async def _refine_one_iterative_async(
        self,
        para: Dict,
        score_info: Dict,
        score_threshold: float,
        max_rounds: int,
        refine_semaphore: asyncio.Semaphore,
        detect_semaphore: asyncio.Semaphore,
        client,
        detect_fn: Callable[[str], Dict],
    ) -> Dict:
        """单段循环润色：润色一次就复测一次，达标或到达轮数上限即停止。"""
        text = para['text']
        aigc_score = score_info.get('aigc_score')

        if aigc_score is not None and aigc_score <= score_threshold:
            return {
                'index': para['index'],
                'chapter': para.get('chapter', 0),
                'original': text,
                'refined': text,
                'strategy': 'skip',
                'changed': False,
                'note': f'AIGC={aigc_score*100:.1f}%，低于阈值，跳过',
                'aigc_before': aigc_score,
                'aigc_after': aigc_score,
                'aigc_after_label': score_info.get('label'),
                'rounds': 0,
                'round_history': [],
            }

        strategy = classify_paragraph(text)
        if strategy == 'skip':
            return {
                'index': para['index'],
                'chapter': para.get('chapter', 0),
                'original': text,
                'refined': text,
                'strategy': 'skip',
                'changed': False,
                'note': '跳过（代码/公式/过短）',
                'aigc_before': aigc_score,
                'aigc_after': aigc_score,
                'aigc_after_label': score_info.get('label'),
                'rounds': 0,
                'round_history': [],
            }

        current_text = text
        current_score = aigc_score
        best_text = text
        best_score = aigc_score
        best_label = score_info.get('label')
        best_note = ""
        history = []

        for round_no in range(1, max_rounds + 1):
            if round_no == 1:
                prompt = PROMPTS[strategy].format(paragraph=current_text)
            else:
                score_value = current_score if current_score is not None else (score_threshold + 0.01)
                prompt = PROMPT_ITERATIVE_HINT_CN.format(
                    score=score_value * 100,
                    paragraph=current_text,
                )

            try:
                async with refine_semaphore:
                    result = await self._call_llm_async(prompt, client)
            except Exception as e:
                logger.warning("段落 %d 第 %d 轮润色失败: %s", para['index'], round_no, e)
                history.append({
                    'round': round_no,
                    'aigc_score': current_score,
                    'changed': False,
                    'note': f'API 调用失败: {e}',
                })
                break

            if not result:
                history.append({
                    'round': round_no,
                    'aigc_score': current_score,
                    'changed': False,
                    'note': 'API 调用失败，保留上一版',
                })
                break

            refined, note = self._parse_refinement_result(result, strategy if round_no == 1 else 'deai_cn')
            if not refined:
                history.append({
                    'round': round_no,
                    'aigc_score': current_score,
                    'changed': False,
                    'note': '模型返回空文本，保留上一版',
                })
                break

            detect_result = await self._detect_async(refined, detect_fn, detect_semaphore)
            next_score = detect_result.get('aigc_score')
            changed = refined.strip() != current_text.strip()
            history.append({
                'round': round_no,
                'aigc_score': next_score,
                'label': detect_result.get('label'),
                'changed': changed,
                'note': note,
            })

            if next_score is not None and (best_score is None or next_score < best_score):
                best_text = refined
                best_score = next_score
                best_label = detect_result.get('label')
                best_note = note

            current_text = refined
            current_score = next_score

            if next_score is not None and next_score <= score_threshold:
                break

        changed = best_text.strip() != text.strip()
        if history:
            final_note = f"循环润色 {len(history)} 轮"
            if best_score is not None:
                final_note += f"，最佳 AIGC={best_score*100:.1f}%"
            if best_note:
                final_note += f"；{best_note}"
        else:
            final_note = "未执行润色"

        return {
            'index': para['index'],
            'chapter': para.get('chapter', 0),
            'original': text,
            'refined': best_text,
            'strategy': strategy,
            'changed': changed,
            'note': final_note,
            'aigc_before': aigc_score,
            'aigc_after': best_score,
            'aigc_after_label': best_label,
            'rounds': len(history),
            'round_history': history,
        }

    async def refine_batch_iterative_async(
        self,
        paragraphs: List[Dict],
        aigc_scores: List[Dict],
        detect_fn: Callable[[str], Dict],
        score_threshold: float = 0.4,
        max_rounds: int = 2,
        detect_concurrency: Optional[int] = None,
        show_progress: bool = True,
    ) -> List[Dict]:
        """并发执行“单段润色 -> 单段检测 -> 循环调整”的降重工作流。"""
        client = self._get_async_client()
        refine_semaphore = asyncio.Semaphore(self.concurrency)
        detect_semaphore = asyncio.Semaphore(detect_concurrency or self.concurrency)
        start_time = time.time()

        tasks = [
            asyncio.create_task(self._with_index(
                i,
                self._refine_one_iterative_async(
                    para,
                    score_info,
                    score_threshold,
                    max(1, max_rounds),
                    refine_semaphore,
                    detect_semaphore,
                    client,
                    detect_fn,
                ),
            ))
            for i, (para, score_info) in enumerate(zip(paragraphs, aigc_scores))
        ]
        results = [None] * len(tasks)
        completed = changed_count = rounds = 0
        bar_ctx = progress_bar(total=len(tasks), desc="循环润色") if show_progress else None
        try:
            if bar_ctx:
                with bar_ctx as bar:
                    for task in asyncio.as_completed(tasks):
                        i, result = await task
                        results[i] = result
                        completed += 1
                        if isinstance(result, dict):
                            changed_count += int(result.get('changed', False))
                            rounds += result.get('rounds', 0)
                            bar.set_postfix(changed=changed_count, rounds=rounds)
                        bar.update(1)
            else:
                for task in asyncio.as_completed(tasks):
                    i, result = await task
                    results[i] = result
                    completed += 1
        except Exception:
            for task in tasks:
                if not task.done():
                    task.cancel()
            raise
        await client.close()

        final_results = []
        for i, r in enumerate(results):
            if isinstance(r, Exception):
                logger.error("段落 %d 循环润色异常: %s", i, r)
                final_results.append({
                    'index': paragraphs[i]['index'],
                    'chapter': paragraphs[i].get('chapter', 0),
                    'original': paragraphs[i]['text'],
                    'refined': paragraphs[i]['text'],
                    'strategy': 'error',
                    'changed': False,
                    'note': f'异常: {r}',
                    'aigc_before': aigc_scores[i].get('aigc_score'),
                    'aigc_after': aigc_scores[i].get('aigc_score'),
                    'aigc_after_label': aigc_scores[i].get('label'),
                    'rounds': 0,
                    'round_history': [],
                })
            else:
                final_results.append(r)

        elapsed = time.time() - start_time
        changed_count = sum(1 for r in final_results if r['changed'])
        rounds = sum(r.get('rounds', 0) for r in final_results)
        logger.info("循环润色完成: %d 段中修改 %d 段，总轮次 %d，耗时 %.1f 秒 (润色并发=%d, 检测并发=%d)",
                    len(final_results), changed_count, rounds, elapsed,
                    self.concurrency, detect_concurrency or self.concurrency)
        return final_results

    async def refine_batch_async(
        self,
        paragraphs: List[Dict],
        aigc_scores: List[Dict],
        score_threshold: float = 0.4,
        show_progress: bool = True,
    ) -> List[Dict]:
        """并发批量润色高风险段落"""
        client = self._get_async_client()
        semaphore = asyncio.Semaphore(self.concurrency)
        start_time = time.time()

        tasks = [
            asyncio.create_task(self._with_index(
                i,
                self._refine_one_async(para, score_info, score_threshold, semaphore, client),
            ))
            for i, (para, score_info) in enumerate(zip(paragraphs, aigc_scores))
        ]
        results = [None] * len(tasks)
        completed = changed_count = skipped_count = 0
        bar_ctx = progress_bar(total=len(tasks), desc="并发润色") if show_progress else None
        try:
            if bar_ctx:
                with bar_ctx as bar:
                    for task in asyncio.as_completed(tasks):
                        i, result = await task
                        results[i] = result
                        completed += 1
                        if isinstance(result, dict):
                            changed_count += int(result.get('changed', False))
                            skipped_count += int(not result.get('changed', False) and result.get('strategy') == 'skip')
                            bar.set_postfix(changed=changed_count, skipped=skipped_count)
                        bar.update(1)
            else:
                for task in asyncio.as_completed(tasks):
                    i, result = await task
                    results[i] = result
                    completed += 1
        except Exception:
            for task in tasks:
                if not task.done():
                    task.cancel()
            raise

        await client.close()

        # 处理异常结果
        final_results = []
        for i, r in enumerate(results):
            if isinstance(r, Exception):
                logger.error("段落 %d 润色异常: %s", i, r)
                final_results.append({
                    'index': paragraphs[i]['index'],
                    'chapter': paragraphs[i].get('chapter', 0),
                    'original': paragraphs[i]['text'],
                    'refined': paragraphs[i]['text'],
                    'strategy': 'error',
                    'changed': False,
                    'note': f'异常: {r}',
                    'aigc_before': aigc_scores[i].get('aigc_score'),
                    'aigc_after': None,
                })
            else:
                final_results.append(r)

        elapsed = time.time() - start_time
        changed_count = sum(1 for r in final_results if r['changed'])
        skipped_count = sum(1 for r in final_results if not r['changed'] and r['strategy'] == 'skip')
        logger.info("并发润色完成: %d 段中修改 %d 段，跳过 %d 段，耗时 %.1f 秒 (并发数=%d)",
                     len(final_results), changed_count, skipped_count, elapsed, self.concurrency)
        return final_results

    def refine_batch(
        self,
        paragraphs: List[Dict],
        aigc_scores: List[Dict],
        score_threshold: float = 0.4,
    ) -> List[Dict]:
        """
        批量润色高风险段落

        Args:
            paragraphs: 段落列表 [{text, chapter, section, index, ...}, ...]
            aigc_scores: AIGC 检测结果 [{aigc_score, label, ...}, ...]
            score_threshold: 润色阈值，分数 > 此值的段落会被润色

        Returns:
            润色结果列表 [{index, original, refined, strategy, changed, note, ...}, ...]
        """
        results = []
        total = len(paragraphs)
        refined_count = 0
        skipped_count = 0
        start_time = time.time()

        for i, (para, score_info) in enumerate(zip(paragraphs, aigc_scores)):
            text = para['text']
            aigc_score = score_info.get('aigc_score')

            # 低于阈值的段落跳过
            if aigc_score is not None and aigc_score <= score_threshold:
                results.append({
                    'index': para['index'],
                    'chapter': para.get('chapter', 0),
                    'original': text,
                    'refined': text,
                    'strategy': 'skip',
                    'changed': False,
                    'note': f'AIGC={aigc_score*100:.1f}%，低于阈值，跳过',
                    'aigc_before': aigc_score,
                    'aigc_after': None,
                })
                skipped_count += 1
                continue

            # 分类并润色
            strategy = classify_paragraph(text)

            if strategy == 'skip':
                results.append({
                    'index': para['index'],
                    'chapter': para.get('chapter', 0),
                    'original': text,
                    'refined': text,
                    'strategy': 'skip',
                    'changed': False,
                    'note': '跳过（代码/公式/过短）',
                    'aigc_before': aigc_score,
                    'aigc_after': None,
                })
                skipped_count += 1
                continue

            refined, note = self.refine_paragraph(text, strategy)
            changed = (refined.strip() != text.strip())
            if changed:
                refined_count += 1

            results.append({
                'index': para['index'],
                'chapter': para.get('chapter', 0),
                'original': text,
                'refined': refined,
                'strategy': strategy,
                'changed': changed,
                'note': note,
                'aigc_before': aigc_score,
                'aigc_after': None,  # 稍后复测
            })

            # 进度
            if (i + 1) % 10 == 0 or i == 0 or i == total - 1:
                elapsed = time.time() - start_time
                logger.info("  润色进度: %d/%d (已修改%d段, 跳过%d段, %.0f秒)",
                            i + 1, total, refined_count, skipped_count, elapsed)

        elapsed = time.time() - start_time
        logger.info("润色完成: %d 段中修改 %d 段，跳过 %d 段，耗时 %.1f 秒",
                     total, refined_count, skipped_count, elapsed)
        return results


# ═══════════════════════════════════════════════════════════
# 报告生成
# ═══════════════════════════════════════════════════════════

def generate_refinement_report(
    refinement_results: List[Dict],
    output_path: str,
    aigc_before: Optional[Dict] = None,
    aigc_after: Optional[Dict] = None,
):
    """生成润色报告"""
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

    changed = [r for r in refinement_results if r['changed']]
    skipped = [r for r in refinement_results if not r['changed']]

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("                    降 AIGC 率润色报告\n")
        f.write("=" * 80 + "\n\n")
        f.write(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"总段落数: {len(refinement_results)}\n")
        f.write(f"已修改:   {len(changed)}\n")
        f.write(f"已跳过:   {len(skipped)}\n\n")

        if aigc_before and aigc_after:
            f.write("-" * 80 + "\n")
            f.write("AIGC 率对比:\n")
            f.write(f"  润色前: {aigc_before.get('ai_ratio', 'N/A')}%\n")
            f.write(f"  润色后: {aigc_after.get('ai_ratio', 'N/A')}%\n")
            delta = aigc_before.get('ai_ratio', 0) - aigc_after.get('ai_ratio', 0)
            f.write(f"  降幅:   {delta:.1f} 个百分点\n")
            f.write("-" * 80 + "\n\n")

        # 策略统计
        strategy_counts = {}
        for r in refinement_results:
            s = r.get('strategy', 'unknown')
            strategy_counts[s] = strategy_counts.get(s, 0) + 1

        f.write("策略分布:\n")
        strategy_labels = {
            'refine_cn': '表达润色（中文）',
            'deai_cn': '去AI化（中文）',
            'deai_en': '去AI化（英文）',
            'logic_check_en': '逻辑检查（英文）',
            'skip': '跳过',
        }
        for s, count in sorted(strategy_counts.items(), key=lambda x: -x[1]):
            label = strategy_labels.get(s, s)
            f.write(f"  {label}: {count} 段\n")
        f.write("\n")

        # 逐段详情
        f.write("-" * 80 + "\n")
        f.write("逐段详情:\n")
        f.write("-" * 80 + "\n\n")

        for r in refinement_results:
            idx = r['index']
            ch = r.get('chapter', '?')
            strategy = r.get('strategy', '?')
            label = strategy_labels.get(strategy, strategy)
            score_before = r.get('aigc_before')
            score_str = f"AIGC={score_before*100:.1f}%" if score_before else "N/A"

            f.write(f"[段落 {idx:3d}] 第{ch}章 | {label} | {score_str}\n")

            if r['changed']:
                f.write(f"  原文: {r['original'][:200]}\n")
                f.write(f"  改后: {r['refined'][:200]}\n")
                f.write(f"  说明: {r['note']}\n")
            else:
                f.write(f"  原文: {r['original'][:200]}\n")
                f.write(f"  状态: {r['note']}\n")

            round_history = r.get('round_history') or []
            if round_history:
                f.write("  循环记录:")
                for item in round_history:
                    score = item.get('aigc_score')
                    score_text = f"{score*100:.1f}%" if score is not None else "N/A"
                    f.write(f" 第{item.get('round')}轮={score_text}")
                f.write("\n")

            f.write("\n")

    logger.info("润色报告已保存: %s", output_path)
    return output_path


def export_refined_docx(
    paragraphs: List[Dict],
    refinement_results: List[Dict],
    output_path: str,
):
    """将润色后的文本导出为 DOCX"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        logger.error("需要安装 python-docx")
        return None

    # 建立 index → refined 映射
    refined_map = {}
    for r in refinement_results:
        if r['changed']:
            refined_map[r['index']] = r['refined']

    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('降 AIGC 率润色后文本（供对照替换）')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}  |  已修改 {len(refined_map)} 段')
    run.font.size = Pt(10.5)

    doc.add_paragraph()

    # 逐段输出
    for para in paragraphs:
        idx = para['index']
        text = refined_map.get(idx, para['text'])

        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0.74)
        pf.line_spacing = Pt(22)

        if idx in refined_map:
            # 修改过的段落：标注 [已润色]
            run = p.add_run(f'[已润色] ')
            run.font.size = Pt(10.5)
            run.font.color.rgb = None  # 默认色
            run.bold = True
            run2 = p.add_run(text)
            run2.font.size = Pt(12)
            run2.font.name = '宋体'
            run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    logger.info("润色后 DOCX 已保存: %s", output_path)
    return output_path
